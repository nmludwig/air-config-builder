"""
Generate a formatted Word doc matching the Dr. Liu AIR Configuration Playbook style.
python-docx — pure Python, no Node.js required.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re, io
from datetime import datetime

# ── Colors ────────────────────────────────────────────────
NAVY    = RGBColor(0x1E,0x3A,0x5F)
BLUE    = RGBColor(0x0F,0x62,0xFE)
GREEN   = RGBColor(0x16,0x65,0x34)
RED     = RGBColor(0x99,0x1B,0x1B)
AMBER   = RGBColor(0x92,0x40,0x0E)
SLATE   = RGBColor(0x47,0x55,0x69)
GRAY4   = RGBColor(0x94,0xA3,0xB8)
TEXT2   = RGBColor(0x33,0x41,0x55)
WHITE   = RGBColor(0xFF,0xFF,0xFF)

NAVY_H  = "1E3A5F"
BLUE_H  = "0F62FE"
BLUE_LT = "EFF6FF"
BLUE_MD = "DBEAFE"
BLUE_BD = "BFDBFE"
GRN_LT  = "DCFCE7"
GRN_BD  = "BBF7D0"
GREEN_DK= "166534"
RED_LT  = "FEE2E2"
RED_BD  = "FECACA"
AMB_LT  = "FEF3C7"
AMB_BD  = "FDE68A"
GRAY_H  = "F3F4F6"
GRAY2_H = "E2E8F0"
WHITE_H = "FFFFFF"

# ── XML Helpers ───────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for s in tcPr.findall(qn('w:shd')): tcPr.remove(s)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="CCCCCC", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for b in tcPr.findall(qn('w:tcBorders')): tcPr.remove(b)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def no_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for b in tcPr.findall(qn('w:tcBorders')): tcPr.remove(b)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tcBorders.append(b)
    tcPr.append(tcBorders)

def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for m in tcPr.findall(qn('w:tcMar')): tcPr.remove(m)
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top',top),('bottom',bottom),('left',left),('right',right)]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)

def set_col_width(cell, w):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for x in tcPr.findall(qn('w:tcW')): tcPr.remove(x)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(w))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def set_table_width(table, w=9360):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for x in tblPr.findall(qn('w:tblW')): tblPr.remove(x)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(w))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

def pspacing(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    for s in pPr.findall(qn('w:spacing')): pPr.remove(s)
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), str(before))
    sp.set(qn('w:after'), str(after))
    pPr.append(sp)

def para_border_bottom(para, color=BLUE_H, sz=12):
    pPr = para._p.get_or_add_pPr()
    for pb in pPr.findall(qn('w:pBdr')): pPr.remove(pb)
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), str(sz))
    b.set(qn('w:space'), '4')
    b.set(qn('w:color'), color)
    pBdr.append(b)
    pPr.append(pBdr)

def clean(text):
    """Strip markdown bold/italic markers"""
    t = re.sub(r'\*\*(.+?)\*\*', r'\1', str(text))
    t = re.sub(r'\*(.+?)\*', r'\1', t)
    return t.strip()

# ── Run helper ────────────────────────────────────────────
def run(para, text, bold=False, italic=False, color=None, size=10):
    r = para.add_run(str(text))
    r.bold = bold
    r.italic = italic
    r.font.name = 'Arial'
    r.font.size = Pt(size)
    if color: r.font.color.rgb = color
    return r

# ── Block builders ────────────────────────────────────────
def add_h1(doc, text):
    p = doc.add_paragraph()
    pspacing(p, 320, 140)
    para_border_bottom(p, BLUE_H, 12)
    run(p, clean(text), bold=True, color=NAVY, size=18)

def add_h2(doc, text):
    p = doc.add_paragraph()
    pspacing(p, 220, 100)
    run(p, clean(text), bold=True, color=NAVY, size=14)

def add_h3(doc, text):
    p = doc.add_paragraph()
    pspacing(p, 160, 80)
    run(p, clean(text), bold=True, color=BLUE, size=12)

def add_body(doc, text, bold=False, italic=False, color=None, size=10):
    p = doc.add_paragraph()
    pspacing(p, 40, 40)
    run(p, clean(text), bold=bold, italic=italic, color=color or TEXT2, size=size)
    return p

def add_label_value(doc, label, value):
    p = doc.add_paragraph()
    pspacing(p, 40, 40)
    run(p, clean(label) + ': ', bold=True, color=NAVY, size=10)
    run(p, clean(value), size=10)

def add_bullet(doc, text, bold_pre=None):
    p = doc.add_paragraph(style='List Bullet')
    pspacing(p, 30, 30)
    if bold_pre:
        run(p, clean(bold_pre) + ': ', bold=True, color=NAVY, size=10)
        run(p, clean(text), size=10)
    else:
        run(p, clean(text), size=10)

def add_numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    pspacing(p, 30, 30)
    run(p, clean(text), size=10)

def add_spacer(doc, pts=8):
    p = doc.add_paragraph()
    pspacing(p, 0, 0)
    p.paragraph_format.line_spacing = Pt(pts)

def add_copy_box(doc, text):
    """Green-tinted monospace box — company description"""
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, 9360)
    cell = table.cell(0,0)
    set_col_width(cell, 9360)
    set_cell_bg(cell, "F0FDF4")
    set_cell_borders(cell, "86EFAC", "4")
    set_cell_margins(cell, 160, 160, 200, 200)
    p = cell.paragraphs[0]
    pspacing(p, 0, 0)
    r = p.add_run(clean(text))
    r.font.name = 'Courier New'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x16,0x65,0x34)
    add_spacer(doc)

def add_dark_box(doc, text):
    """Dark navy box with white italic text — for greetings"""
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, 9360)
    cell = table.cell(0,0)
    set_col_width(cell, 9360)
    set_cell_bg(cell, NAVY_H)
    set_cell_borders(cell, "2D4E78", "4")
    set_cell_margins(cell, 160, 160, 200, 200)
    p = cell.paragraphs[0]
    pspacing(p, 0, 0)
    # Handle BUSINESS HOURS / AFTER-HOURS labels inside box
    text = clean(text)
    for label in ["BUSINESS HOURS GREETING:", "AFTER-HOURS GREETING:"]:
        if label in text:
            parts = text.split(label)
            for j, part in enumerate(parts):
                if j > 0:
                    lp = cell.add_paragraph()
                    pspacing(lp, 60, 0)
                    lr = lp.add_run(label)
                    lr.bold = True
                    lr.font.name = 'Arial'
                    lr.font.size = Pt(10)
                    lr.font.color.rgb = WHITE
                    if part.strip():
                        cp = cell.add_paragraph()
                        pspacing(cp, 0, 60)
                        cr = cp.add_run(part.strip())
                        cr.italic = True
                        cr.font.name = 'Arial'
                        cr.font.size = Pt(10)
                        cr.font.color.rgb = WHITE
            return
    r = p.add_run(text)
    r.italic = True
    r.font.name = 'Arial'
    r.font.size = Pt(10)
    r.font.color.rgb = WHITE
    add_spacer(doc)

def add_callout(doc, icon, label, body_text, bg, text_color):
    table = doc.add_table(rows=1, cols=2)
    set_table_width(table, 9360)
    c0, c1 = table.rows[0].cells[0], table.rows[0].cells[1]
    set_col_width(c0, 480); no_borders(c0); set_cell_bg(c0, bg)
    set_cell_margins(c0, 120, 120, 120, 80)
    p0 = c0.paragraphs[0]; pspacing(p0, 0, 0)
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p0, icon, size=12, color=text_color)
    set_col_width(c1, 8880); no_borders(c1); set_cell_bg(c1, bg)
    set_cell_margins(c1, 120, 120, 100, 140)
    p1 = c1.paragraphs[0]; pspacing(p1, 0, 0)
    if label:
        run(p1, clean(label) + ':  ', bold=True, color=text_color, size=10)
    run(p1, clean(body_text), color=text_color, size=10)
    add_spacer(doc)

def info_box(doc, label, text):
    add_callout(doc, 'i', label, text, BLUE_MD, RGBColor(0x1E,0x3A,0x8A))

def warn_box(doc, label, text):
    add_callout(doc, '!', label, text, AMB_LT, AMBER)

def danger_box(doc, label, text):
    add_callout(doc, '!', label, text, RED_LT, RED)

def success_box(doc, label, text):
    add_callout(doc, '✓', label, text, GRN_LT, RGBColor(0x16,0x65,0x34))

def add_faq_table(doc, title, question_variants, answer):
    if title:
        add_h3(doc, title)
    C1, C2 = 3400, 5960
    table = doc.add_table(rows=2, cols=2)
    set_table_width(table, 9360)
    hr = table.rows[0]
    for cell, txt, w in [(hr.cells[0],"QUESTION (what caller says)",C1),(hr.cells[1],"ANSWER (what AIR says — copy exactly)",C2)]:
        set_col_width(cell, w); set_cell_bg(cell, NAVY_H); set_cell_borders(cell, NAVY_H)
        set_cell_margins(cell, 80,80,120,80)
        p = cell.paragraphs[0]; pspacing(p,0,0)
        run(p, txt, bold=True, color=WHITE, size=9)
    cr = table.rows[1]
    c0, c1 = cr.cells[0], cr.cells[1]
    set_col_width(c0, C1); set_col_width(c1, C2)
    set_cell_bg(c0, BLUE_LT); set_cell_borders(c0, GRAY2_H)
    set_cell_bg(c1, GRAY_H);  set_cell_borders(c1, GRAY2_H)
    set_cell_margins(c0, 100,100,120,80); set_cell_margins(c1, 100,100,120,120)
    p0 = c0.paragraphs[0]; pspacing(p0,0,0)
    run(p0, clean(question_variants), italic=True, size=9)
    p1 = c1.paragraphs[0]; pspacing(p1,0,0)
    run(p1, clean(answer), size=9)
    add_spacer(doc)

def add_routing_table(doc, rules):
    if not rules: return
    C = [3200, 2160, 4000]
    table = doc.add_table(rows=1+len(rules), cols=3)
    set_table_width(table, sum(C))
    for j, (hdr, w) in enumerate(zip(["TRIGGER KEYWORDS / PHRASES","ROUTE TO","COVERS / NOTES"], C)):
        cell = table.rows[0].cells[j]
        set_col_width(cell, w); set_cell_bg(cell, NAVY_H); set_cell_borders(cell, NAVY_H)
        set_cell_margins(cell, 80,80,120,80)
        p = cell.paragraphs[0]; pspacing(p,0,0)
        run(p, hdr, bold=True, color=WHITE, size=9)
    for i, rule in enumerate(rules):
        bg1 = BLUE_LT if i%2==0 else WHITE_H
        bg2 = GRAY_H  if i%2==0 else WHITE_H
        row = table.rows[i+1]
        for j, (val, bg, bld, ital) in enumerate(zip(
            [rule.get('keywords',''), rule.get('route',''), rule.get('note','')],
            [bg1, bg2, bg2], [False,True,False], [True,False,False])):
            cell = row.cells[j]
            set_col_width(cell, C[j]); set_cell_bg(cell, bg); set_cell_borders(cell, GRAY2_H)
            set_cell_margins(cell, 80,80,120,80)
            p = cell.paragraphs[0]; pspacing(p,0,0)
            run(p, clean(val), bold=bld, italic=ital, size=9)
    add_spacer(doc)

def add_kv_table(doc, rows, header_row=False):
    if not rows: return
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_width(table, 9360)
    for i, (k, v) in enumerate(rows):
        bg_l = BLUE_LT if (header_row and i==0) else GRAY_H
        bg_r = BLUE_MD if (header_row and i==0) else (GRAY_H if i%2==0 else WHITE_H)
        c0, c1 = table.rows[i].cells
        set_col_width(c0, 2800); set_col_width(c1, 6560)
        set_cell_bg(c0, bg_l); set_cell_borders(c0, GRAY2_H)
        set_cell_bg(c1, bg_r); set_cell_borders(c1, GRAY2_H)
        set_cell_margins(c0, 80,80,120,80); set_cell_margins(c1, 80,80,120,120)
        p0 = c0.paragraphs[0]; pspacing(p0,0,0)
        run(p0, clean(k), bold=True, color=NAVY if not (header_row and i==0) else WHITE, size=10)
        p1 = c1.paragraphs[0]; pspacing(p1,0,0)
        run(p1, clean(v), bold=(header_row and i==0), color=WHITE if (header_row and i==0) else None, size=10)
    add_spacer(doc)

def add_coverage_table(doc, rows):
    C = [360, 2800, 2400, 3800]
    table = doc.add_table(rows=1+len(rows), cols=4)
    set_table_width(table, sum(C))
    for j, (hdr, w) in enumerate(zip(["#","Call Reason","AIR Handles?","How"], C)):
        cell = table.rows[0].cells[j]
        set_col_width(cell, w); set_cell_bg(cell, NAVY_H); set_cell_borders(cell, NAVY_H)
        set_cell_margins(cell, 80,80,120,80)
        p = cell.paragraphs[0]; pspacing(p,0,0)
        run(p, hdr, bold=True, color=WHITE, size=9)
    for i, row_data in enumerate(rows):
        num, reason, handles, how = row_data[0], row_data[1], row_data[2], row_data[3]
        h_upper = clean(handles).upper()
        if 'YES' in h_upper:   h_bg, h_col = "1A5C2B", WHITE
        elif 'PARTIAL' in h_upper: h_bg, h_col = "7C4A00", WHITE
        else:                       h_bg, h_col = "8B1A1A", WHITE
        bg = GRAY_H if i%2==0 else WHITE_H
        row = table.rows[i+1]
        vals  = [str(num), clean(reason), clean(handles), clean(how)]
        bgs   = [bg, bg, h_bg, bg]
        cols  = [None, None, RGBColor(0xFF,0xFF,0xFF), None]
        bolds = [False, False, True, False]
        aligns= [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT]
        for j, (cell, val, bg_, col_, bld, aln) in enumerate(zip(row.cells, vals, bgs, cols, bolds, aligns)):
            set_col_width(cell, C[j]); set_cell_bg(cell, bg_); set_cell_borders(cell, GRAY2_H)
            set_cell_margins(cell, 80,80,120,80)
            p = cell.paragraphs[0]; pspacing(p,0,0); p.alignment = aln
            run(p, val, bold=bld, color=col_, size=9)
    add_spacer(doc)

# ── Table parser ──────────────────────────────────────────
def parse_md_table(lines, start_idx):
    rows = []
    i = start_idx
    while i < len(lines):
        s = lines[i].strip()
        if not s.startswith('|'):
            break
        if re.match(r'^\|[\s\-|]+\|$', s):
            i += 1; continue
        cells = [c.strip() for c in s.strip('|').split('|')]
        cells = [re.sub(r'\*\*(.+?)\*\*', r'\1', c) for c in cells]
        cells = [re.sub(r'\*(.+?)\*', r'\1', c) for c in cells]
        rows.append([c.strip() for c in cells])
        i += 1
    return rows, i

def render_table(doc, rows):
    if not rows: return
    headers = rows[0]
    data = rows[1:]
    h0 = headers[0].upper() if headers else ''

    # Coverage table
    if len(headers) == 4 and headers[0].strip() in ('#','') or (len(headers) == 4 and 'AIR' in headers[2].upper()):
        coverage = [(r[0],r[1],r[2],r[3]) for r in data if len(r)==4]
        if coverage: add_coverage_table(doc, coverage); return

    # FAQ table
    if len(headers) == 2 and 'QUESTION' in h0:
        for r in data:
            if len(r) == 2: add_faq_table(doc, '', r[0], r[1])
        return

    # Routing table
    if len(headers) == 3 and any(k in h0 for k in ['TRIGGER','KEYWORD']):
        rules = [{'keywords':r[0],'route':r[1],'note':r[2]} for r in data if len(r)==3]
        if rules: add_routing_table(doc, rules); return

    # Single-cell table — treat as copy/dark box
    if len(headers) == 1 and not data:
        add_copy_box(doc, headers[0]); return

    # 2-col KV
    if len(headers) == 2:
        all_rows = rows  # include header as first row
        add_kv_table(doc, [(r[0],r[1]) for r in all_rows if len(r)==2]); return

    # Fallback
    for r in data:
        if r: add_body(doc, ' | '.join(r))

# ── Main renderer ─────────────────────────────────────────
def render(doc, content):
    lines = content.split('\n')
    n = len(lines)
    i = 0
    route_rules = []
    section_count = 0

    def flush_routing():
        nonlocal route_rules
        if route_rules:
            add_routing_table(doc, route_rules)
            route_rules = []

    while i < n:
        line = lines[i]
        s = line.strip()

        if not s:
            i += 1; continue

        # ── Major section H1: # N. Title or **N. Title** ──
        m1 = re.match(r'^#\s+(\d+)\.\s+(.+)', s)
        m1b = re.match(r'^\*\*(\d+)\.\s+(.+?)\*\*$', s)
        if m1 or m1b:
            flush_routing()
            m = m1 or m1b
            num, title = int(m.group(1)), m.group(2).replace('**','')
            if section_count > 0:
                doc.add_page_break()
            section_count += 1
            add_h1(doc, f"{num}. {title}")
            i += 1; continue

        # ── H2: ## N.N Title ──────────────────────────────
        m2 = re.match(r'^#{2,3}\s+(\d+\.\d+)\s+(.+)', s)
        if m2:
            flush_routing()
            add_spacer(doc, 6)
            add_h2(doc, f"{m2.group(1)} {m2.group(2).replace('**','')}")
            i += 1; continue

        # ── H3: ### Call Reason N ─────────────────────────
        m3 = re.match(r'^#{3,4}\s+(.+)', s)
        if m3:
            text = m3.group(1).replace('**','')
            add_h3(doc, text)
            i += 1; continue

        # ── Bold section header acting as H2 ─────────────
        # e.g. **2. AIR Setup** or **3. Skills Configuration**
        mbs = re.match(r'^\*\*(\d+)\.\s+(.+?)\*\*$', s)
        if mbs:
            flush_routing()
            num, title = int(mbs.group(1)), mbs.group(2)
            if section_count > 0:
                doc.add_page_break()
            section_count += 1
            add_h1(doc, f"{num}. {title}")
            i += 1; continue

        # ── Bold subsection: **2.1 Title** ───────────────
        mbss = re.match(r'^\*\*(\d+\.\d+)\s+(.+?)\*\*$', s)
        if mbss:
            flush_routing()
            add_spacer(doc, 6)
            add_h2(doc, f"{mbss.group(1)} {mbss.group(2)}")
            i += 1; continue

        # ── FAQ heading: **Call Reason N — Topic** ────────
        mfaq = re.match(r'^\*\*(Call Reason\s+\d+[^*]+)\*\*$', s, re.I)
        if mfaq:
            flush_routing()
            add_h3(doc, mfaq.group(1))
            i += 1; continue

        # ── Blockquote callouts: > ℹ️ or > ⚠️ or > ✅ ────
        if s.startswith('> '):
            text = s[2:].strip()
            emoji = text[0] if text else ''
            rest = text[1:].strip() if text else ''
            # Parse "Label: body" or just body
            label_m = re.match(r'\*\*(.+?):\*\*\s*(.*)', rest)
            if emoji in ('ℹ️', 'ℹ'):
                if label_m: info_box(doc, label_m.group(1), label_m.group(2))
                else: info_box(doc, 'Note', clean(rest))
            elif emoji in ('⚠️', '⚠'):
                if label_m: warn_box(doc, label_m.group(1), label_m.group(2))
                else: warn_box(doc, 'Note', clean(rest))
            elif emoji in ('✅', '✓'):
                if label_m: success_box(doc, label_m.group(1), label_m.group(2))
                else: success_box(doc, 'Note', clean(rest))
            elif emoji in ('🔴', '❌'):
                if label_m: danger_box(doc, label_m.group(1), label_m.group(2))
                else: danger_box(doc, 'Note', clean(rest))
            else:
                add_body(doc, clean(rest))
            i += 1; continue

        # ── Inline callouts: i | text or ! | text ────────
        m_callout = re.match(r'^([i!✓⚠️❌✅])\s*\|\s*(.+)', s)
        if m_callout:
            icon, text = m_callout.group(1), m_callout.group(2)
            label_m = re.match(r'\*\*(.+?):\*\*\s*(.*)', text)
            lbl = label_m.group(1) if label_m else ''
            bod = label_m.group(2) if label_m else text
            if icon == 'i': info_box(doc, lbl, bod)
            elif icon in ('!', '⚠️'): warn_box(doc, lbl, bod)
            elif icon in ('✓', '✅'): success_box(doc, lbl, bod)
            elif icon in ('❌', '🔴'): danger_box(doc, lbl, bod)
            i += 1; continue

        # ── Emoji callouts: ⚠️ text / ✅ text ─────────────
        if s[:2] in ('⚠️','✅','❌','🔴') or s[:1] in ('⚠','✅','❌'):
            emoji = s[:2] if s[:2] in ('⚠️','✅','❌','🔴') else s[:1]
            text = s[len(emoji):].strip()
            label_m = re.match(r'\*\*(.+?):\*\*\s*(.*)', text)
            lbl = label_m.group(1) if label_m else ''
            bod = label_m.group(2) if label_m else text
            if '⚠' in emoji: warn_box(doc, lbl, bod)
            elif '✅' in emoji or '✓' in emoji: success_box(doc, lbl, bod)
            elif '❌' in emoji or '🔴' in emoji: danger_box(doc, lbl, bod)
            i += 1; continue

        # ── Code blocks ───────────────────────────────────
        if s.startswith('```'):
            code = []
            i += 1
            while i < n and not lines[i].strip().startswith('```'):
                code.append(lines[i])
                i += 1
            i += 1
            t = '\n'.join(code).strip()
            if t: add_copy_box(doc, t)
            continue

        # ── Markdown tables ───────────────────────────────
        if s.startswith('|'):
            tbl_rows, new_i = parse_md_table(lines, i)
            if tbl_rows:
                render_table(doc, tbl_rows)
            i = new_i
            continue

        # ── Horizontal rule ───────────────────────────────
        if re.match(r'^-{3,}$', s):
            i += 1; continue

        # ── Transfer by Context rule block ────────────────
        rule_m = re.match(r'^\*\*Rule\s*\d*[:\s]*(.+?)\*\*', s)
        if rule_m:
            rule = {'title': rule_m.group(1), 'keywords':'', 'route':'', 'note':''}
            i += 1
            while i < n:
                ls = lines[i].strip()
                km = re.match(r'^Keywords?[:\s]+(.+)', ls, re.I)
                rm = re.match(r'^Route\s+to[:\s]+(.+)', ls, re.I)
                nm = re.match(r'^(?:Covers|Note|AIR|Holding)[:\s]+(.+)', ls, re.I)
                if km: rule['keywords'] = clean(km.group(1)); i+=1; continue
                if rm: rule['route'] = clean(rm.group(1)); i+=1; continue
                if nm: rule['note'] = clean(nm.group(1)).strip('"'); i+=1; continue
                if re.match(r'^\*\*Rule|^#{2,}|^---$', ls): break
                i += 1
            route_rules.append(rule)
            continue

        # ── Bullets ───────────────────────────────────────
        if re.match(r'^[-*]\s+', s):
            text = re.sub(r'^[-*]\s+', '', s)
            bm = re.match(r'\*\*(.+?)\*\*[:\s—-]+(.+)', text)
            if bm: add_bullet(doc, clean(bm.group(2)), clean(bm.group(1)))
            else: add_bullet(doc, clean(text))
            i += 1; continue

        # ── Numbered list ─────────────────────────────────
        if re.match(r'^\d+\.\s+', s):
            add_numbered(doc, clean(re.sub(r'^\d+\.\s+', '', s)))
            i += 1; continue

        # ── Bold label: value ─────────────────────────────
        bkv = re.match(r'^\*\*(.+?)\*\*::\s*(.+)', s)
        if not bkv:
            bkv = re.match(r'^\*\*(.+?)\*\*[:\s]+(.+)', s)
        if bkv:
            k, v = clean(bkv.group(1)), clean(bkv.group(2))
            add_label_value(doc, k, v)
            i += 1; continue

        # ── Bold only line ────────────────────────────────
        if re.match(r'^\*\*.+\*\*$', s):
            add_body(doc, clean(s), bold=True, color=NAVY)
            i += 1; continue

        # ── Admin portal path lines ───────────────────────
        if '→' in s and ('Admin Portal' in s or 'Phone System' in s):
            add_body(doc, clean(s), italic=True, color=GRAY4, size=9)
            i += 1; continue

        # ── Plain text ────────────────────────────────────
        text = clean(s)
        if text:
            add_body(doc, text)
        i += 1

    flush_routing()


# ── Main ──────────────────────────────────────────────────
def generate_docx(biz_name, content, prepared_by="RingCentral SE"):
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(1)
        section.left_margin = section.right_margin = Inches(1)

    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(10)

    today = datetime.now().strftime('%B %Y')
    biz_name = re.sub(r'\*+', '', biz_name).strip()

    # ── Cover ──────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pspacing(p, 480, 80)
    run(p, "RingCentral AI Receptionist", bold=True, color=NAVY, size=26)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pspacing(p, 0, 80)
    run(p, "Configuration Playbook", color=BLUE, size=20)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pspacing(p, 80, 40)
    run(p, biz_name, bold=True, color=SLATE, size=16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pspacing(p, 20, 20)
    run(p, "Ready-to-paste configuration for every AIR field", italic=True, color=GRAY4, size=10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pspacing(p, 20, 20)
    run(p, f"Prepared by: {prepared_by}  |  {today}", italic=True, color=GRAY4, size=9)

    doc.add_page_break()

    render(doc, content)

    # ── Footer ──────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pspacing(p, 240, 60)
    para_border_bottom(p, GRAY2_H, 4)
    run(p, f"RingCentral AI Receptionist  |  {biz_name}  |  {today}",
        italic=True, color=GRAY4, size=8)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
